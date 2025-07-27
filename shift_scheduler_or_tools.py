import streamlit as st
import pandas as pd
from collections import defaultdict, OrderedDict
from ortools.sat.python import cp_model
from docx import Document
import io
import random

st.title("Smart Shift Scheduler (OR-Tools, Gender & Tech-Aware)")

# === 1. INPUT SECTION ===
st.subheader("Upload Excel File (Name, Gender (M/F), Technical (Yes/No))")
excel_file = st.file_uploader("Upload Excel File", type=[".xlsx"])

names_data = []
if excel_file:
    df = pd.read_excel(excel_file)
    for _, row in df.iterrows():
        name = str(row[0]).strip()
        gender = str(row[1]).strip().upper()
        tech = str(row[2]).strip().lower() == "yes"
        if name:
            names_data.append((name, gender, tech))

# === 2. CATEGORIZE PEOPLE ===
males = [n for n, g, _ in names_data if g == 'M']
females = [n for n, g, _ in names_data if g == 'F']
tech_males = [n for n, g, t in names_data if g == 'M' and t]
tech_females = [n for n, g, t in names_data if g == 'F' and t]
nontech_males = list(set(males) - set(tech_males))
nontech_females = list(set(females) - set(tech_females))

# === 3. TEXT INPUTS FOR POSTS ===
common_male = st.text_area("COMMON Duty Posts for Males (priority)").split("\n")
common_female = st.text_area("COMMON Duty Posts for Females (priority)").split("\n")
c_only = st.text_area("C-Shift ONLY Posts (Male Only)").split("\n")
gen_male = st.text_area("General Shift Posts (Male)").split("\n")
gen_female = st.text_area("General Shift Posts (Female)").split("\n")
tech_posts = st.text_area("TECHNICAL General Shift Posts").split("\n")
merge_input = st.text_area("MERGE Groups (comma-separated, e.g., Gate1, Gate2)").split("\n")

common_male = [p.strip() for p in common_male if p.strip()]
common_female = [p.strip() for p in common_female if p.strip()]
c_only = [p.strip() for p in c_only if p.strip()]
gen_male = [p.strip() for p in gen_male if p.strip()]
gen_female = [p.strip() for p in gen_female if p.strip()]
tech_posts = [p.strip() for p in tech_posts if p.strip()]
merge_groups = [[s.strip() for s in line.split(",") if s.strip()] for line in merge_input if line.strip()]

# === 4. PREVIOUS SHIFTS AND HISTORY ===
st.subheader("Enter Previous Day Shift Info (Name: Shift)")
prev_input = st.text_area("e.g., John: C")
prev_shift_map = {}
if prev_input.strip():
    for line in prev_input.strip().split("\n"):
        if ":" in line:
            name, shift = line.split(":")
            prev_shift_map[name.strip()] = shift.strip()

st.subheader("Upload Weekly History")
hist_file = st.file_uploader("Weekly Shift History Excel", type=[".xlsx"], key="weekly_hist")
weekly_counts = defaultdict(lambda: {"C": 0, "Night12": 0, "Day12": 0})
if hist_file:
    hist_df = pd.read_excel(hist_file)
    for _, row in hist_df.iterrows():
        name = row[0]
        weekly_counts[name] = {
            "C": int(row.get("C", 0)),
            "Night12": int(row.get("Night12", 0)),
            "Day12": int(row.get("Day12", 0))
        }

# === 5. DUTY POST PLANNING (with merge & 12hr conversion) ===
def plan_duty_posts():
    post_plan = OrderedDict()
    shift_demand = defaultdict(int)

    # 1. Technical Posts
    for post in tech_posts:
        post_plan[post] = ["GS"]
        shift_demand["GS"] += 1

    # 2. General Posts
    for post in gen_male:
        post_plan[post] = ["GS"]
        shift_demand["GS"] += 1
    for post in gen_female:
        post_plan[post] = ["GS"]
        shift_demand["GS"] += 1

    # 3. C-only Posts (males only)
    for post in c_only:
        post_plan[post] = ["C"]
        shift_demand["C"] += 1

    # 4. Female common posts
    for post in common_female:
        post_plan[post] = ["A", "B", "C"]
        shift_demand["A"] += 1
        shift_demand["B"] += 1
        shift_demand["C"] += 1  # will assign male to C

    # 5. Male common posts
    merged_map = {}
    merged_set = set()
    for group in merge_groups:
        if len(group) > 1:
            top = group[0]
            for p in group[1:]:
                merged_map[p] = top
                merged_set.add(p)

    reduced = []
    seen = set()
    for post in common_male:
        if post in merged_set or post in seen:
            continue
        reduced.append(post)
        seen.add(post)

    post_status = OrderedDict((p, ["A", "B", "C"]) for p in reduced)

    # Now convert to 12hr from bottom if not enough people
    total_needed = len(post_status) * 3 + sum(len(v) for k, v in post_plan.items())
    while total_needed > len(males) + len(females):
        for post in reversed(reduced):
            if post_status[post] == ["A", "B", "C"]:
                post_status[post] = ["Day12", "Night12"]
                break
        total_needed = sum(len(v) for v in post_status.values()) + sum(len(v) for v in post_plan.values())

    for post, shifts in post_status.items():
        post_plan[post] = shifts
        for s in shifts:
            shift_demand[s] += 1

    return post_plan, shift_demand

# === 6. OR-TOOLS SOLVER ===
def solve_with_ortools(post_plan, shift_demand):
    people = [n for n, _, _ in names_data]
    model = cp_model.CpModel()
    shifts = ["A", "B", "C", "Day12", "Night12", "GS", "Off"]
    shift_vars = {(p, s): model.NewBoolVar(f"{p}_{s}") for p in people for s in shifts}

    for p in people:
        model.AddExactlyOne(shift_vars[p, s] for s in shifts)

    # Assign shift demands
    for s in shift_demand:
        model.Add(sum(shift_vars[p, s] for p in people) == shift_demand[s])

    # No A/Day12/GS after Night12 or C
    for p in people:
        if prev_shift_map.get(p) in ["C", "Night12"]:
            model.Add(shift_vars[p, "A"] == 0)
            model.Add(shift_vars[p, "Day12"] == 0)
            model.Add(shift_vars[p, "GS"] == 0)

    # Weekly shift limits
    flags = []
    max_4_allowed = max(1, int(0.1 * len(people)))
    for p in people:
        wc = weekly_counts[p]
        total_night = model.NewIntVar(0, 10, f"tn_{p}")
        total_12hr = model.NewIntVar(0, 10, f"t12_{p}")
        model.Add(total_night == wc["C"] + wc["Night12"] + shift_vars[p,"C"])
        model.Add(total_12hr == wc["Day12"] + wc["Night12"] + shift_vars[p,"Day12"] + shift_vars[p,"Night12"])
        model.Add(total_12hr <= 3)

        flag = model.NewBoolVar(f"flag_{p}")
        gt3 = model.NewBoolVar(f"gt3_{p}")
        model.Add(total_night > 3).OnlyEnforceIf(gt3)
        model.Add(total_night <= 3).OnlyEnforceIf(gt3.Not())
        model.Add(flag == 1).OnlyEnforceIf(gt3)
        model.Add(flag == 0).OnlyEnforceIf(gt3.Not())
        flags.append(flag)
    model.Add(sum(flags) <= max_4_allowed)

    solver = cp_model.CpSolver()
    solver.parameters.random_seed = random.randint(1, 10000)
    status = solver.Solve(model)
    return status, solver, shift_vars

# === 7. GENERATE OUTPUT ===
if st.button("Generate Schedule"):
    post_plan, shift_demand = plan_duty_posts()
    status, solver, shift_vars = solve_with_ortools(post_plan, shift_demand)

    if status == cp_model.OPTIMAL or status == cp_model.FEASIBLE:
        assignment = {}
        for p in [n for n, _, _ in names_data]:
            for s in ["A", "B", "C", "Day12", "Night12", "GS", "Off"]:
                if solver.Value(shift_vars[p, s]):
                    assignment[p] = s

        # Assign people to posts
        shift_buckets = defaultdict(list)
        for p, s in assignment.items():
            shift_buckets[s].append(p)
        for s in shift_buckets:
            random.shuffle(shift_buckets[s])

        doc = Document()
        doc.add_heading("Shift Schedule", 0)

        for post, shifts in post_plan.items():
            doc.add_paragraph(post, style='Heading 2')
            for s in shifts:
                person = shift_buckets[s].pop() if shift_buckets[s] else "Unfilled"
                doc.add_paragraph(f"{s}: {person}", style='List Bullet')

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        st.download_button("Download Shift Schedule", buffer, file_name="shift_schedule.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    else:
        st.error("❌ Could not generate a feasible schedule. Please check constraints or staff availability.")
